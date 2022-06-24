# -*- coding: utf-8 -*-
# Author: Vladimir Pichugin <vladimir@pichug.in>
import os

from docx import Document
from docx2pdf import convert

from utils.helpers import L10n
from utils.data import Contract

from settings import Settings


class ContractGenerator:
    @staticmethod
    def create_contract(file_name, contract: Contract):
        program = contract.get_program()
        created_dt = contract.created_dt()

        programType = L10n.get('program.type.distance') if 'distance' in program else L10n.get('program.type.fulltime')
        programName = L10n.get('program.{program}.code'.format(program=program)) + ' ' + L10n.get('program.{program}.name'.format(program=program))
        programDuration = L10n.get('program.{program}.duration'.format(program=program))

        customerDocWhenIssuedDt = contract.get_dt('customer_docWhenIssued')
        customerDocWhenIssued = '{day} {month} {year}'.format(
            day=customerDocWhenIssuedDt.strftime('%d'),
            month=L10n.get('months.gen.{month}'.format(month=customerDocWhenIssuedDt.strftime('%B'))),
            year=customerDocWhenIssuedDt.strftime('%Y')
        )

        studentDocWhenIssuedDt = contract.get_dt('student_docWhenIssued')
        studentDocWhenIssued = ''
        if studentDocWhenIssuedDt:
            studentDocWhenIssued = '{day} {month} {year}'.format(
                day=studentDocWhenIssuedDt.strftime('%d'),
                month=L10n.get('months.gen.{month}'.format(month=studentDocWhenIssuedDt.strftime('%B'))),
                year=studentDocWhenIssuedDt.strftime('%Y')
            )

        studentSex = L10n.get('sex.female') if contract.get_sex('student') == 'female' else L10n.get('sex.male')
        representativeSex = L10n.get('sex.female') if contract.get_sex('customer') == 'female' else L10n.get('sex.male')

        placeholders = {
            'contractId': contract.get_contract_id(),
            'emailConfirmCode': contract.get_confirm_code(),

            'programType': programType,
            'programName': programName,
            'programDuration': programDuration,

            'd': created_dt.strftime('%d'),
            'm': L10n.get('months.gen.{month}'.format(month=created_dt.strftime('%B'))),
            'y': created_dt.strftime('%Y'),

            'programCR': L10n.get('program.{program}.cr'.format(program=program)),
            'programCRW': L10n.get('program.{program}.crw'.format(program=program)),
            'programDurationAccelerated': '_______ года ______ месяцев',
            'programCYR': L10n.get('program.{program}.cyr'.format(program=program)),
            'programCYRW': L10n.get('program.{program}.cyrw'.format(program=program)),

            'representative': contract.get_name('customer', True),
            'representativeSex': representativeSex,
            'representativeDocSerial': contract.get('customer_docSerial'),
            'representativeDocNumber': contract.get('customer_docNumber'),
            'representativeDocWhenIssued': customerDocWhenIssued,
            'representativeDocIssued': contract.get('customer_docIssued'),
            'representativeAddress': contract.get('customer_address'),
            'representativePhone': contract.get('customer_phone'),
            'representativeEmail': contract.get('customer_email'),

            'company': L10n.get('company'),
            'companyShort': L10n.get('companyShort'),
            'companyLicenseId': L10n.get('companyLicenseId'),
            'accreditation': L10n.get('accreditation'),

            'director': L10n.get('director'),
            'directorAfterSignature': L10n.get('directorAfterSignature')
        }

        customer_role = contract.get_customer_role()
        customer_age = contract.get_customer_age()
        if customer_role == '1' and customer_age >= 18:
            student = {
                'student': '',
                'studentSex': '',
                'studentDocSerial': '',
                'studentDocNumber': '',
                'studentDocWhenIssued': '',
                'studentDocIssued': '',
                'studentAddress': '',
                'studentPhone': '',
                'studentEmail': ''
            }
        else:
            student = {
                'student': contract.get_name('student', True),
                'studentSex': studentSex,
                'studentDocSerial': contract.get('student_docSerial'),
                'studentDocNumber': contract.get('student_docNumber'),
                'studentDocWhenIssued': studentDocWhenIssued,
                'studentDocIssued': contract.get('student_docIssued'),
                'studentAddress': contract.get('student_address'),
                'studentPhone': contract.get('student_phone'),
                'studentEmail': contract.get('student_email')
            }
        placeholders.update(student)

        doc = Document(os.path.join(os.getcwd(), Settings.CONTRACT_TEMPLATE))

        ContractGenerator.docx_replace(doc, placeholders)

        file_path = os.path.join(Settings.SAVE_CONTRACTS_FOLDER, f'{file_name}.docx')

        doc.save(file_path)
        convert(file_path)

        return True

    @staticmethod
    def docx_replace(doc, data):
        paragraphs = list(doc.paragraphs)

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)

        for p in paragraphs:
            for key, val in data.items():
                key_name = '${{{}}}'.format(key)  # Формат плейсхолдера: ${PlaceholderName}

                if key_name in p.text:
                    inline = p.runs

                    # Заменит строки и сохранит прежний стиль.
                    # Заменяемый текст может быть разделен на несколько прогонов, поэтому
                    # выполнит поиск, определит, какие прогоны нуждаются в замене текста,
                    # а затем заменит текст в тех, которые были определены.
                    started = False
                    key_index = 0

                    # found_runs это лист (inline index, index of match, length of match)
                    found_runs = list()
                    found_all = False
                    replace_done = False
                    for i in range(len(inline)):
                        # Случай №1: найден в одиночном прогоне, поэтому замкните накоротко замену
                        if key_name in inline[i].text and not started:
                            found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                            text = inline[i].text.replace(key_name, str(val))
                            inline[i].text = text
                            replace_done = True
                            found_all = True
                            break

                        if key_name[key_index] not in inline[i].text and not started:
                            continue  # Продолжаем искать...

                        # Случай №2: Поиск частичного текста, поиск первого запуска.
                        if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                            # Проверяем последовательность.
                            start_index = inline[i].text.find(key_name[key_index])
                            check_length = len(inline[i].text)
                            for text_index in range(start_index, check_length):
                                if inline[i].text[text_index] != key_name[key_index]:
                                    break  # Нет совпадений.
                            if key_index == 0:
                                started = True
                            chars_found = check_length - start_index
                            key_index += chars_found
                            found_runs.append((i, start_index, chars_found))
                            if key_index != len(key_name):
                                continue
                            else:
                                # Нашёл все символы в key_name.
                                found_all = True
                                break

                        # Случай №2: Поиск частичного текста, поиск последующего запуска.
                        if key_name[key_index] in inline[i].text and started and not found_all:
                            # Проверяем последовательность.
                            chars_found = 0
                            check_length = len(inline[i].text)
                            for text_index in range(0, check_length):
                                try:
                                    if inline[i].text[text_index] == key_name[key_index]:
                                        key_index += 1
                                        chars_found += 1
                                    else:
                                        break
                                except IndexError:
                                    pass
                                    #print(text_index)
                                    #print(key_index)

                            # Не совпадает, должен быть это конец.
                            found_runs.append((i, 0, chars_found))
                            if key_index == len(key_name):
                                found_all = True
                                break

                    if found_all and not replace_done:
                        for i, item in enumerate(found_runs):
                            index, start, length = [t for t in item]
                            if i == 0:
                                text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                                inline[index].text = text
                            else:
                                text = inline[index].text.replace(inline[index].text[start:start + length], '')
                                inline[index].text = text
                    # print(p.text)
