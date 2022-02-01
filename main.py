from docx import Document
from docx2pdf import convert


def docx_replace(doc, data):
	paragraphs = list(doc.paragraphs)

	for t in doc.tables:
		for row in t.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					paragraphs.append(paragraph)

	for p in paragraphs:
		for key, val in data.items():
			key_name = '${{{}}}'.format(key)  # Placeholders format: ${PlaceholderName}

			if key_name in p.text:
				inline = p.runs

				# Replace strings and retain the same style.
				# The text to be replaced can be split over several runs so
				# search through, identify which runs need to have text replaced
				# then replace the text in those identified
				started = False
				key_index = 0

				# found_runs is a list of (inline index, index of match, length of match)
				found_runs = list()
				found_all = False
				replace_done = False
				for i in range(len(inline)):
					# case 1: found in single run so short circuit the replace
					if key_name in inline[i].text and not started:
						found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
						text = inline[i].text.replace(key_name, str(val))
						inline[i].text = text
						replace_done = True
						found_all = True
						break

					if key_name[key_index] not in inline[i].text and not started:
						continue  # keep looking ...

					# case 2: search for partial text, find first run
					if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
						# check sequence
						start_index = inline[i].text.find(key_name[key_index])
						check_length = len(inline[i].text)
						for text_index in range(start_index, check_length):
							if inline[i].text[text_index] != key_name[key_index]:
								break  # no match so must be false positive
						if key_index == 0:
							started = True
						chars_found = check_length - start_index
						key_index += chars_found
						found_runs.append((i, start_index, chars_found))
						if key_index != len(key_name):
							continue
						else:
							# found all chars in key_name
							found_all = True
							break

					# case 2: search for partial text, find subsequent run
					if key_name[key_index] in inline[i].text and started and not found_all:
						# check sequence
						chars_found = 0
						check_length = len(inline[i].text)
						for text_index in range(0, check_length):
							try:
								if inline[i].text[text_index] == key_name[key_index]:
									key_index += 1
									chars_found += 1
								else:
									break
							except IndexError:
								print(text_index)
								print(key_index)
						# no match so must be end
						found_runs.append((i, 0, chars_found))
						if key_index == len(key_name):
							found_all = True
							break

				if found_all and not replace_done:
					for i, item in enumerate(found_runs):
						index, start, length = [t for t in item]
						if i == 0:
							text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
							inline[index].text = text
						else:
							text = inline[index].text.replace(inline[index].text[start:start + length], '')
							inline[index].text = text
				# print(p.text)


placeholders = {
	'company': 'Государственное бюджетное  профессиональное образовательное учреждение города Москвы «Московский колледж бизнес – технологий»',
	'companyShort': 'ГБПОУ КБТ',
	'companyLicenseId': 'лицензии № 035560 от «28» октября 2014 года',
	'accreditation': 'свидетельства о государственной аккредитации серии 77А01 № 0003589, регистрационный номер 003589, выданного Департаментом образования города Москвы «29» апреля 2015 года',

	'director': 'Аверьяновой Ларисы Васильевны',
	'directorAfterSignature': 'Л.В. Аверьянова',

	'contractId': 'test-0/2/2022',
	'emailConfirmCode': 'test',

	'd': '1',
	'm': 'февраля',
	'y': '2022',
	'representative': 'Иванова Мария Ивановна',
	'representativeSex': 'ая',
	'student': 'Иванов Владимир Иванович',
	'studentSex': 'ый',

	'programName': '09.02.07 Информационные системы и программирование',
	'programDuration': '3 года 10 месяцев',
	'programDurationAccelerated': '_______ года ______ месяцев',

	'representativeDocSerial': '00 00',
	'representativeDocNumber': '000 000',
	'representativeDocWhenIssued': '1 января 2022',
	'representativeDocIssued': 'ГУ МВД РОССИИ ПО ГОРОДУ МОСКВЕ',
	'representativeAddress': 'Москва',
	'representativePhone': '+7 777 777 77 77',
	'representativeEmail': 'mail@internet.ru',

	'studentDocSerial': '00 00',
	'studentDocNumber': '000 000',
	'studentDocWhenIssued': '1 января 2022',
	'studentDocIssued': 'ГУ МВД РОССИИ ПО ГОРОДУ МОСКВЕ',
	'studentAddress': 'Москва',
	'studentPhone': '+7 777 777 77 77',
	'studentEmail': 'mail@internet.ru',

	'programCR': '390 000',
	'programCRW': 'триста девяносто тысяч рублей',

	'programCYR': '130 000',
	'programCYRW': 'триста девяносто тысяч рублей',

}

doc = Document('templates/contract_ochnoe.docx')
docx_replace(doc, placeholders)
doc.save('temp.docx')
convert('temp.docx')
